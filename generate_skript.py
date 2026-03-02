# -*- coding: utf-8 -*-
"""
Generiert LFA_IHK_Praesentationsskript.docx
1:1 Spiegelung der Landing-Page-Reihenfolge
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── Farben ─────────────────────────────────────────────────────────────────
RED      = RGBColor(0xE0, 0x1B, 0x24)   # LFA-Rot
DARK     = RGBColor(0x0D, 0x0D, 0x1A)   # Dunkel
WALERI   = RGBColor(0x7C, 0x3A, 0xED)   # Violett  → Waleri (Gründer)
DANIEL   = RGBColor(0x06, 0x7A, 0xFF)   # Blau     → Daniel (Ausbilder)
LEON_C   = RGBColor(0x05, 0x96, 0x69)   # Grün     → Leon (Azubi)
ELIAS_C  = RGBColor(0xD9, 0x77, 0x06)   # Orange   → Elias (Azubi)
STAGE    = RGBColor(0x6B, 0x72, 0x80)   # Grau     → Regieanweisungen
BG_BOX   = RGBColor(0xF3, 0xF4, 0xF6)   # Hellgrau → Infoboxen

SPEAKER_COLORS = {
    "WALERI": WALERI,
    "DANIEL": DANIEL,
    "LEON":   LEON_C,
    "ELIAS":  ELIAS_C,
}

SPEAKER_LABELS = {
    "WALERI": "Waleri Moretz  |  Gründer & Ausbilder",
    "DANIEL": "Daniel Moretz  |  Ausbilder & Fachlicher Leiter",
    "LEON":   "Leon Moretz  |  FIAE-Auszubildender",
    "ELIAS":  "Elias  |  FIAE-Auszubildender",
}

# ── Hilfsfunktionen ────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Setzt Hintergrundfarbe einer Tabellenzelle (OOXML)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_table_border(table):
    """Entfernt alle Rahmen einer Tabelle (randlos)."""
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),  "none")
        border.set(qn("w:sz"),   "0")
        border.set(qn("w:space"),"0")
        border.set(qn("w:color"),"auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def speaker_block(doc, speaker_key: str, text: str, timing: str = ""):
    """
    Erzeugt einen farbigen Sprecher-Block:
    [Farbbalken | Label + Timing]
    [Weißes Feld | Gesprochener Text    ]
    """
    color     = SPEAKER_COLORS[speaker_key]
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    label     = SPEAKER_LABELS[speaker_key]

    # ── Kopfzeile (Farbbalken) ──
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.autofit = False
    header_tbl.columns[0].width = Cm(0.6)
    header_tbl.columns[1].width = Cm(15.4)
    add_table_border(header_tbl)

    cell_bar   = header_tbl.cell(0, 0)
    cell_label = header_tbl.cell(0, 1)
    set_cell_bg(cell_bar, hex_color)
    set_cell_bg(cell_label, "F8F9FA")

    p_label = cell_label.paragraphs[0]
    p_label.clear()
    run = p_label.add_run(f"  {label}")
    run.bold = True
    run.font.size  = Pt(8.5)
    run.font.color.rgb = color
    if timing:
        run2 = p_label.add_run(f"    ⏱ {timing}")
        run2.font.size  = Pt(7.5)
        run2.font.color.rgb = STAGE
    p_label.paragraph_format.space_before = Pt(0)
    p_label.paragraph_format.space_after  = Pt(0)

    # ── Text-Bereich ──
    body_tbl = doc.add_table(rows=1, cols=2)
    body_tbl.autofit = False
    body_tbl.columns[0].width = Cm(0.6)
    body_tbl.columns[1].width = Cm(15.4)
    add_table_border(body_tbl)

    cell_side = body_tbl.cell(0, 0)
    cell_text = body_tbl.cell(0, 1)
    set_cell_bg(cell_side, hex_color)
    set_cell_bg(cell_text, "FFFFFF")

    p_text = cell_text.paragraphs[0]
    p_text.clear()
    run = p_text.add_run(f"  {text}")
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    p_text.paragraph_format.space_before = Pt(2)
    p_text.paragraph_format.space_after  = Pt(8)
    p_text.paragraph_format.left_indent  = Cm(0.3)

    doc.add_paragraph()   # Abstand nach Block


def stage_direction(doc, text: str):
    """Kursive Regieanweisung in Grau."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"[ {text} ]")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = STAGE


def section_heading(doc, number: str, title: str, subtitle: str = ""):
    """Abschnittsüberschrift mit roter Nummerierung."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    run_num   = p.add_run(f"{number}  ")
    run_num.bold = True
    run_num.font.size = Pt(16)
    run_num.font.color.rgb = RED
    run_title = p.add_run(title.upper())
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(2)
    if subtitle:
        ps = doc.add_paragraph(subtitle)
        ps.runs[0].font.size = Pt(9)
        ps.runs[0].font.color.rgb = STAGE
        ps.runs[0].italic = True
        ps.paragraph_format.space_before = Pt(0)
        ps.paragraph_format.space_after  = Pt(6)
    # Trennlinie
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(0)
    p_hr.paragraph_format.space_after  = Pt(10)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E01B24")
    pBdr.append(bottom)
    pPr.append(pBdr)


def info_box(doc, text: str):
    """Hellgrauer Info-Kasten für Zahlen / Schlüsselinfos."""
    tbl = doc.add_table(rows=1, cols=1)
    add_table_border(tbl)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, "EFF6FF")
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    run.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    doc.add_paragraph()


def transition_note(doc, text: str):
    """Übergangsnotiz zwischen Sprechern."""
    p = doc.add_paragraph()
    run = p.add_run(f"→  {text}")
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)


# ── Hauptdokument ──────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Seitenränder
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.2)
        sec.right_margin  = Cm(2.5)

    # Standardabsatz-Schriftart
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ══════════════════════════════════════════════════════════════════════
    # TITELSEITE
    # ══════════════════════════════════════════════════════════════════════
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(60)
    r = p_title.add_run("LFA")
    r.bold = True
    r.font.size  = Pt(52)
    r.font.color.rgb = RED

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_sub.add_run("Die KI-gestützte Plattform für die duale IT-Ausbildung")
    r2.font.size = Pt(14)
    r2.font.color.rgb = DARK

    doc.add_paragraph()
    p_line = doc.add_paragraph("─" * 62)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.runs[0].font.color.rgb = RED

    doc.add_paragraph()
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p_meta.add_run(
        "Präsentationsskript  ·  IHK Frankfurt am Main\n"
        "Vertraulich  ·  WAMOCON GmbH  ·  2026"
    )
    r3.font.size = Pt(10)
    r3.font.color.rgb = STAGE

    doc.add_paragraph()
    doc.add_paragraph()

    # Farb-Legende
    legend_tbl = doc.add_table(rows=1, cols=4)
    add_table_border(legend_tbl)
    for idx, (key, color) in enumerate(SPEAKER_COLORS.items()):
        hex_c = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        cell  = legend_tbl.cell(0, idx)
        set_cell_bg(cell, hex_c)
        p_l = cell.paragraphs[0]
        p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_l = p_l.add_run(SPEAKER_LABELS[key].split("  |  ")[0])
        r_l.bold = True
        r_l.font.size = Pt(9)
        r_l.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p_legend = doc.add_paragraph()
    p_legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p_legend.add_run("Farb-Legende der Sprecher")
    rl.font.size = Pt(8)
    rl.italic = True
    rl.font.color.rgb = STAGE

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # VORBEMERKUNGEN
    # ══════════════════════════════════════════════════════════════════════
    p_h = doc.add_paragraph()
    r_h = p_h.add_run("Hinweise zur Nutzung dieses Skripts")
    r_h.bold = True
    r_h.font.size = Pt(13)
    r_h.font.color.rgb = DARK
    p_h.paragraph_format.space_after = Pt(8)

    hinweise = [
        "Dieses Skript ist wörtlich ausformuliert — ihr müsst es nicht ablesen, sondern nutzt es als Vorlage und formuliert in euren eigenen Worten.",
        "Regieanweisungen stehen in grau-kursiven Klammern [ … ] — diese werden NICHT gesprochen.",
        "Alle Sprecher sind namentlich farblich markiert. Stimmt euch vorher ab, wer welchen Teil übernimmt.",
        "Gesamtdauer: ca. 25–30 Minuten + 10 Minuten Fragerunde.",
        "Die Reihenfolge folgt exakt der Landing Page — scrollt gemeinsam durch die Seite während ihr sprecht.",
        "Empfehlung: Generalprobe mindestens einmal, idealerweise mit Live-Demo auf fiae-learn.com/demo.",
    ]
    for h in hinweise:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(h).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()
    info_box(doc,
        "⏱  Timing-Überblick:  Einstieg 2 Min  ·  Problem 3 Min  ·  Lösung & Rollen 3 Min  ·  "
        "Warum LFA 3 Min  ·  Funktionen 2 Min  ·  Demo 4 Min  ·  Tätigkeitsnachweis 3 Min  ·  "
        "ROI 2 Min  ·  Architektur 2 Min  ·  6 Gründe 3 Min  ·  Abschluss & Fragen 5–8 Min"
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # 0  VOR DER SEITE — Einstieg  (kein Bildschirm)
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "0", "Einstieg — Vor der Landing Page",
                    "Noch kein Bildschirm · Vertrauen aufbauen · ~2 Minuten")

    stage_direction(doc, "Alle Beteiligten stehen vor dem Publikum. Bildschirm ist noch schwarz oder zeigt nur das WAMOCON-Logo.")

    speaker_block(doc, "WALERI",
        "Schön, dass das heute geklappt hat — danke, dass ihr euch die Zeit nehmt. "
        "Ich bin Waleri, Gründer der WAMOCON GmbH und gleichzeitig Ausbilder. "
        "Mit mir dabei: Daniel, mein Mitausbilder und fachlicher Leiter — "
        "sowie Leon und Elias, unsere beiden Auszubildenden, die seit August 2025 dabei sind "
        "und aktiv an dem mitgearbeitet haben, was ich euch heute zeigen will.\n\n"
        "Wir machen heute keine Folien-Präsentation. Wir zeigen euch ein echtes, laufendes Produkt — "
        "entwickelt in Frankfurt, gebaut für die IT-Ausbildung, getestet in unserem eigenen Betrieb. "
        "Es heißt LFA. Das Lernzentrum für Auszubildende.\n\n"
        "Kurze Vorbemerkung: Wir haben das nicht gebaut, weil wir ein Startup gründen wollten. "
        "Sondern weil wir das Problem aus jeder Perspektive kennen — als ehemaliger Azubi, als Ausbilder, und als Betrieb. "
        "Und weil wir gemerkt haben, dass es dafür noch keine brauchbare Lösung gibt.",
        "~2 Min")

    transition_note(doc, "Bildschirm einschalten → Browser öffnen → Landing Page laden: lfa-learn.com (oder lokal)")

    # ══════════════════════════════════════════════════════════════════════
    # 1  HERO
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "1", "Hero — Lernzentrum für Auszubildende",
                    "Landing Page ist jetzt sichtbar · Erster Eindruck · ~1 Minute")

    stage_direction(doc, "Landing Page lädt, Hero-Animation läuft — Typing-Animation 'FIAE Ausbildung' / 'IT-Karriere' ist sichtbar.")

    speaker_block(doc, "WALERI",
        "Was Sie gerade sehen, ist LFA — das Lernzentrum für Auszubildende. "
        "Eine Plattform, die mit Hilfe von künstlicher Intelligenz speziell für die IT-Ausbildung entwickelt wurde. "
        "Nicht für Schulen. Nicht für Universitäten. Für Betriebe, die IT-Fachkräfte ausbilden — "
        "genau wie wir.\n\n"
        "Die Auszeichnung oben links — 'KI Innovator — Deutsches Innovationsinstitut' — haben wir nicht selbst vergeben. "
        "Das ist eine externe Bewertung unseres Ansatzes.\n\n"
        "Lassen Sie uns jetzt direkt zur Frage kommen, die diese Plattform überhaupt erst nötig gemacht hat.",
        "~1 Min")

    transition_note(doc, "Langsam nach unten scrollen → Problemsektion")

    # ══════════════════════════════════════════════════════════════════════
    # 2  PROBLEMSTELLUNG
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "2", "Problemstellung — Betriebe scheitern an der Organisation",
                    "Drei Statistiken · Drei Problemkarten · ~3 Minuten")

    stage_direction(doc, "Seite scrollt in die Problemsektion. Die drei großen Zahlen sind sichtbar: 109.000 / 76% / 30%.")

    speaker_block(doc, "WALERI",
        "Drei Zahlen. Die erste kennen Sie alle aus dem Bitkom-Report 2025: "
        "109.000 fehlende IT-Fachkräfte in Deutschland — heute, nicht in zehn Jahren. Heute.\n\n"
        "Die zweite: 76 Prozent aller Ausbilder fühlen sich dauerhaft überlastet. "
        "Das ist kein Einzelschicksal — das ist eine strukturelle Erschöpfung im System. "
        "Quelle: DGB-Index Gute Arbeit.\n\n"
        "Und die dritte Zahl ist die, die mich persönlich am meisten trifft: "
        "30 Prozent Abbruchquote bei FIAE. Drei von zehn Azubis kommen nicht durch. "
        "Das sind keine schlechten Schüler. Das ist ein System, das sie nicht hält.",
        "~1:30 Min")

    stage_direction(doc, "Kurze Pause. Dann weiter zu den drei Problemkarten scrollen.")

    info_box(doc,
        "Schlüsselzahlen (Quellen auf der Seite referenzieren):\n"
        "  · 109.000 fehlende IT-Fachkräfte  (Bitkom 2025)\n"
        "  · 76% überlastete Ausbilder  (DGB Index Gute Arbeit)\n"
        "  · 30% Abbruchquote FIAE  (BIBB Datenreport)\n"
        "  · Kosten pro Abbruch: 42.000–49.000 €  (intern kalkuliert)"
    )

    speaker_block(doc, "WALERI",
        "Und wenn man genauer hinschaut, sieht man drei konkrete Ursachen.\n\n"
        "Erstens: Digitaler Papierkram. Die meisten Betriebe haben inzwischen digitale Werkzeuge — "
        "aber die verwalten Dokumente, nicht Wissen und Können. "
        "Das Berichtsheft ist ausgefüllt, aber niemand weiß, ob der Azubi wirklich verstanden hat, was er eingetragen hat.\n\n"
        "Zweitens: Lernen ohne Bezug zur Praxis. Viele Online-Lernplattformen bieten Theorie — "
        "aber IT-Ausbildung ist ein praktischer Beruf. Es gibt kein echtes Lernen ohne Szenario, ohne reale Aufgabe. "
        "Wer nur Lehrbücher ins Digitale überträgt, ändert nichts am Lernfortschritt.\n\n"
        "Drittens: Keine Orientierung. Azubis wissen oft nicht, wo sie gerade stehen, was als nächstes kommt "
        "und was die IHK von ihnen erwartet. Ohne klare Struktur verlieren sie sich im Alltag — "
        "und Ausbilder merken es meistens erst dann, wenn es zu spät ist.",
        "~1:30 Min")

    stage_direction(doc, "Kurze Pause. Dann Überblick auf die 'LFA ändert das'-Bridge.")

    # ══════════════════════════════════════════════════════════════════════
    # 3  LÖSUNG  (id="loesung")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "3", "Lösung — LFA ändert das",
                    "Drei Lösungskarten · Spannungsauflösung · ~2 Minuten")

    stage_direction(doc, "Seite scrollt zur Lösung-Sektion: 'LFA ändert das. Drei Probleme. Eine Plattform.'")

    speaker_block(doc, "DANIEL",
        "Drei Probleme. Eine Plattform. Das ist der Kern von LFA — und ich erkläre kurz, wie wir das konkret umgesetzt haben.\n\n"
        "Problem eins war der Verwaltungsaufwand. Unsere Antwort: Dokumente entstehen von selbst. "
        "Berichte, Bescheinigungen, Zeugnisse — sie werden nicht mehr von Hand geschrieben, "
        "sondern entstehen automatisch aus dem, was der Azubi bereits in der Plattform einträgt. "
        "Der Ausbilder hat wieder Zeit für das, wofür er eigentlich da ist: begleiten und fördern.\n\n"
        "Problem zwei war das Lernen ohne Praxisbezug. Unsere Antwort: Die offiziellen IHK-Bücher werden nicht einfach hochgeladen — "
        "sie werden in echte Aufgaben, Alltagssituationen und Quizze umgewandelt. "
        "Lernen wird greifbar, nicht abstrakt.\n\n"
        "Problem drei war die fehlende Orientierung. Unsere Antwort: Ein klarer Ausbildungsweg. "
        "Lernfortschritt, Kalender, Fristen, Ziele — alles an einem Ort, für Azubi und Ausbilder sichtbar. "
        "Nicht als Kontrolle, sondern als gemeinsamer Kompass.",
        "~2 Min")

    transition_note(doc, "Scrollen zu 'Warum LFA' — die sechs Alleinstellungsmerkmale")

    # ══════════════════════════════════════════════════════════════════════
    # 4  WARUM LFA / SECHS GRÜNDE  (id="roadmap", nav-label "Warum LFA")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "4", "Warum LFA — Sechs Gründe. Eine Plattform.",
                    "Snake-Timeline · Leon & Elias abwechselnd · ~3 Minuten")

    stage_direction(doc, "Roadmap-Sektion sichtbar. Überschrift: 'Sechs Gründe. Eine Plattform.' — Snake-Timeline animiert sich beim Scrollen. Leon und Elias sprechen je drei Punkte.")

    speaker_block(doc, "LEON",
        "Punkt eins — Die KI bleibt auf dem Boden. "
        "Ich habe andere KI-Werkzeuge ausprobiert. Die geben manchmal Antworten, "
        "die gut klingen — aber falsch sind. HAI gibt mir nur Antworten, für die es eine nachweisbare Quelle gibt. "
        "Als Azubi will ich das — weil ich selbst nicht immer weiß, was richtig ist.\n\n"
        "Punkt zwei — Es gibt nichts Vergleichbares. "
        "Wir haben beim Aufbau aktiv danach gesucht: Es gibt kein anderes Produkt auf dem deutschen Markt, "
        "das Lernplattform, Berichtsheft, KI-Lernbegleiter und Ausbilder-Übersicht "
        "so zusammenbringt wie LFA. Das ist kein Versprechen — das ist das Ergebnis unserer eigenen Suche.\n\n"
        "Punkt drei — Ein echter Mensch schaut drauf. "
        "Daniel liest meine Einträge. Er kommentiert. Er benotet. "
        "Das macht einen echten Unterschied — weil ich merke, dass meine Arbeit wahrgenommen wird.",
        "~1:30 Min")

    speaker_block(doc, "ELIAS",
        "Punkt vier — Lernen durch echte Situationen statt trockener Theorie. "
        "In der Schule hatte ich Aufgaben, die ich auswendig gelernt habe — ohne zu verstehen, wofür. "
        "In LFA bekomme ich eine Aufgabe wie aus dem echten Arbeitsalltag: eine Kundenanfrage, "
        "ein Fehler, den ich finden muss, ein Szenario, das ich lösen muss. Das vergisst man nicht.\n\n"
        "Punkt fünf — Das Wissen des Betriebs steckt in der KI. "
        "WAMOCON hat eigene Arbeitsregeln und Abläufe. Die sind in HAI hinterlegt. "
        "Wenn ich frage, wie ich etwas umsetzen soll, "
        "bekomme ich die Antwort, die zu unserem Unternehmen passt — nicht irgendeine allgemeine Antwort.\n\n"
        "Punkt sechs — Alles an einem Ort. "
        "Kein zweites Fenster für das Berichtsheft, kein drittes für den Kalender, kein viertes für die Aufgaben. "
        "LFA ist eine einzige Anlaufstelle. "
        "Das klingt einfach — aber in der Ausbildung ist das alles andere als selbstverständlich.",
        "~1:30 Min")

    transition_note(doc, "Scrollen zu 'Zwei Rollen. Eine Plattform.'")

    # ══════════════════════════════════════════════════════════════════════
    # 5  ROLLEN  (id="rollen")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "5", "Rollen — Zwei Rollen. Eine Plattform.",
                    "Ausbilder- & Azubi-Karte · ~2:30 Minuten")

    stage_direction(doc, "Rollen-Sektion ist sichtbar: Ausbilder-Karte links, Azubi-Karte rechts.")

    speaker_block(doc, "DANIEL",
        "Ich spreche kurz aus der Ausbilder-Perspektive — denn das ist mein Alltag.\n\n"
        "Früher hatte ich keine zentrale Übersicht über meine Azubis. Ich musste fragen, E-Mails schreiben, Dokumente einsammeln. "
        "LFA gibt mir eine Übersichtsseite, auf der ich alle Azubis auf einen Blick sehe: "
        "Wer ist wie weit? Wer hat noch offene Aufgaben? Wer entwickelt sich gut, wer braucht Unterstützung? "
        "Und das alles ohne zusätzlichen Aufwand.\n\n"
        "Dokumente entstehen auf Knopfdruck. Die Plattform ist von Anfang an auf den offiziellen Ausbildungsrahmenplan der IHK abgestimmt. "
        "Und Feedback gebe ich direkt in der Plattform — keine langen E-Mail-Ketten, kein Papierstapel.",
        "~1:15 Min")

    speaker_block(doc, "LEON",
        "Und aus meiner Perspektive als Azubi sieht das so aus:\n\n"
        "Ich weiß jeden Morgen, was von mir erwartet wird. Nicht weil mir jemand eine Liste geschickt hat — "
        "sondern weil mein Kalender, meine offenen Aufgaben und mein Lernfortschritt direkt auf meiner Startseite stehen. "
        "Ich verliere mich nicht im Alltag.\n\n"
        "Was mir persönlich am meisten hilft, ist HAI — unser KI-Lernbegleiter. "
        "Wenn ich bei einer Aufgabe nicht weiterkomme, frage ich HAI. "
        "Er gibt mir keine fertige Lösung — er stellt mir Gegenfragen, "
        "zeigt mir, wo ich nachschauen kann, und hilft mir, selbst auf die Antwort zu kommen. "
        "Das ist kein Abschreib-Werkzeug. Es ist wie ein geduldiger Ansprechpartner, der immer erreichbar ist.",
        "~1:15 Min")

    transition_note(doc, "Scrollen zu 'Funktionen' — Modulübersicht")

    # ══════════════════════════════════════════════════════════════════════
    # 6  FUNKTIONEN  (id="funktionen")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "6", "Funktionen — Alles, was die FIAE-Ausbildung braucht",
                    "8 Module im Überblick · ~2 Minuten")

    stage_direction(doc, "Funktionen-Sektion sichtbar: 2 große Featured-Kacheln (Lernmanagement + HAI) + 6 kleine Modul-Karten.")

    speaker_block(doc, "ELIAS",
        "Acht Bereiche in einer Plattform. Ich führe Sie kurz durch — und hebe zwei heraus, "
        "die ich persönlich täglich nutze.\n\n"
        "Das Herzstück ist die Lernverwaltung: ein klarer Lernpfad, geordnete Aufgaben, "
        "eine Fortschrittsanzeige — alles aufgebaut auf den 12 Lernfeldern der offiziellen FIAE-Ausbildung.\n\n"
        "Daneben: HAI — unser KI-Lernbegleiter. Dazu kommen wir gleich noch ausführlicher.\n\n"
        "Außerdem: ein Quiz-System, das sich dem eigenen Wissensstand anpasst. "
        "Das digitale Berichtsheft — abgestimmt auf die IHK, als PDF exportierbar. "
        "Eine Leistungsbewertung durch den Ausbilder. "
        "Automatisch erstellte Arbeitszeugnisse, auf die jeder Azubi gesetzlich Anspruch hat. "
        "Ein Kalender mit allen wichtigen Terminen. "
        "Und eine Funktion, mit der Ausbilder mehrere Azubis auf einmal verwalten können.\n\n"
        "Das klingt nach viel — weil es viel ist. Aber es fühlt sich nicht überwältigend an, "
        "weil alles miteinander verbunden ist. Kein Bereich steht für sich allein.",
        "~2 Min")

    transition_note(doc, "Scrollen zu 'Tätigkeitsnachweis'")

    # ══════════════════════════════════════════════════════════════════════
    # 7  TÄTIGKEITSNACHWEIS  (id="taetigkeitsnachweis")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "7", "Tätigkeitsnachweis — Zwei Wege, ein Sieger",
                    "Manuell vs. LFA · Zahlen & Belege · ~2:30 Minuten")

    stage_direction(doc, "Tätigkeitsnachweis-Sektion sichtbar. Drei Kennzahlen: 78 Stunden / 156 Dokumente / 156 Datenpunkte.")

    speaker_block(doc, "DANIEL",
        "Ich möchte jetzt auf etwas eingehen, das ihr aus dem Alltag bei der IHK kennt: "
        "das Berichtsheft — oder wie es offiziell heißt: der Tätigkeitsnachweis.\n\n"
        "Die Zahl auf der Seite: 78 Stunden. Das ist die Zeit, die ein Ausbilder "
        "im Durchschnitt über drei Ausbildungsjahre damit verbringt, das Berichtsheft "
        "eines einzigen Azubis zu begleiten und zu prüfen. Bei zehn Azubis: 780 Stunden. "
        "Das sind fast fünf Monate Vollzeit-Arbeit — nur für ein gesetzlich vorgeschriebenes Dokument.\n\n"
        "Mit LFA: unter fünf Minuten pro Azubi und Woche. Der Azubi trägt ein, was er gemacht hat, "
        "und das System ordnet den Eintrag automatisch dem richtigen Ausbildungsbereich zu. "
        "Kein manuelles Prüfen, kein Nachfragen.\n\n"
        "156 Einträge entstehen über drei Jahre — einer pro Woche. "
        "Mit LFA sind das gleichzeitig 156 dokumentierte Belege für das spätere Arbeitszeugnis. "
        "Nicht aus dem Gedächtnis heraus — sondern als lückenloser Nachweis.\n\n"
        "Seit dem 1. August 2024 sind digitale Arbeitszeugnisse in Deutschland rechtlich gültig. "
        "LFA ist von Anfang an darauf ausgelegt.",
        "~2:30 Min")

    info_box(doc,
        "Relevant:  Digitales Arbeitszeugnis rechtlich gültig seit 01.08.2024\n"
        "Berichtsheft: nach offiziellem Ausbildungsrahmenplan gegliedert, als PDF exportierbar"
    )

    transition_note(doc, "Scrollen zum Lernuniversum — Solar System Sektion")

    # ══════════════════════════════════════════════════════════════════════
    # 8  LERNUNIVERSUM  (id="solarsystem")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "8", "Lernuniversum — Vom Chaos zur Ordnung",
                    "Galaxie-Visualisierung · Elias · ~1 Minute · optional")

    stage_direction(doc, "Solar-System-Sektion sichtbar. Überschrift: 'Vom Chaos → zur Ordnung'. Klick auf 'FIAE' öffnet die interaktive Galaxie-Ansicht.")

    speaker_block(doc, "ELIAS",
        "Was ihr jetzt seht, ist kein normales Feature — es ist eher ein Werkzeug zur Orientierung. "
        "Wir nennen es das Lernuniversum.\n\n"
        "Ich klicke auf 'FIAE' — und ihr seht: das gesamte Ausbildungsprogramm, "
        "visualisiert wie eine Galaxie. "
        "12 Lernfelder, aufgeteilt in konkrete Themen, und darunter echte Aufgaben aus dem Arbeitsalltag.\n\n"
        "Das ist der komplette Ausbildungsrahmenplan — nicht als trockene Liste, "
        "sondern so aufgebaut, dass ein Azubi am ersten Tag der Ausbildung sieht, "
        "was ihn in drei Jahren erwartet. Kein Rätselraten mehr.",
        "~1 Min")

    transition_note(doc, "Scrollen zur Demo-Sektion")

    # ══════════════════════════════════════════════════════════════════════
    # 9  DEMO  (id="demo")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "9", "Demo — Live-Plattform",
                    "Echtes Produkt · Leon zeigt beide Ansichten · ~2 Minuten")

    stage_direction(doc, "Demo-Sektion in der Landing Page sichtbar — oder direkt fiae-learn.com/demo öffnen. Unten links in der Demo: Umschalter zwischen Azubi- und Ausbilder-Ansicht.")

    speaker_block(doc, "LEON",
        "Jetzt zeige ich es direkt — genug Worte, hier ist das echte Produkt.\n\n"
        "Ihr seht die Demo von LFA. Ich bin als Azubi angemeldet. "
        "Was ihr seht, ist meine Startseite: meine offenen Aufgaben, mein aktueller Lernfortschritt, "
        "mein nächster Termin. Sofort, ohne erst klicken zu müssen.\n\n"
        "Ich wechsle kurz zum Berichtsheft — das ist das wöchentliche Dokument, "
        "das jeder Azubi laut Gesetz führen muss. Hier trage ich ein, was ich gemacht habe. "
        "Das System ordnet den Eintrag automatisch dem richtigen Ausbildungsbereich zu.\n\n"
        "Jetzt wechsle ich auf die Ausbilder-Ansicht — der Umschalter ist unten links.\n\n"
        "Als Ausbilder sehe ich alle Azubis auf einen Blick, kann meinen letzten Eintrag direkt bewerten "
        "und einen Kommentar hinterlassen. Kein Extra-Aufwand, alles an einem Ort.\n\n"
        "Das ist LFA — nicht als Präsentation, sondern als lauffähiges System.",
        "~2 Min")

    stage_direction(doc, "Demo kurz innehalten — falls Fragen kommen, jetzt eingehen.")

    transition_note(doc, "Scrollen zur HAI-Sektion")

    # ══════════════════════════════════════════════════════════════════════
    # 10  HAI — KI-ASSISTENT  (id="hai")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "10", "HAI — Der KI-Lernbegleiter",
                    "Wie HAI funktioniert · Sokrates-Methode · ~2 Minuten")

    stage_direction(doc, "HAI-Sektion sichtbar. Zeigt den 5-Schritte-Prozess wie HAI eine Antwort erstellt (RAG-Pipeline), einen Chat-Dialog und die Quiz-Progression.")

    speaker_block(doc, "LEON",
        "HAI ist unser KI-Lernbegleiter — und ich möchte kurz erklären, wie er tatsächlich funktioniert. "
        "Das ist wichtig, weil viele denken: KI in der Ausbildung ist riskant.\n\n"
        "Wenn ich eine Frage stelle, sucht HAI nicht im Internet. "
        "Er vergleicht meine Frage mit den Inhalten, die wir ihm freigegeben haben — "
        "IHK-Bücher, Kursmaterial, betriebliche Unterlagen. "
        "Dann antwortet er auf Basis dieser Quellen — und zeigt mir genau, woher die Antwort stammt. "
        "Auf der Seite seht ihr diesen Ablauf Schritt für Schritt erklärt.\n\n"
        "Was HAI nicht tut: fertige Lösungen liefern, aus dem Internet suchen, etwas erfinden.\n\n"
        "Stattdessen: Er stellt mir Gegenfragen. Er verweist mich auf die richtige Stelle. "
        "Er lässt mich selbst draufkommen. Das funktioniert auch beim Quiz — "
        "erst einfache Fragen, dann schwierigere, aber erst wenn ich die vorherigen wirklich verstanden habe.\n\n"
        "Ich lerne nicht, um Punkte zu sammeln. Ich lerne, weil das System mich zwingt, wirklich zu verstehen.",
        "~2 Min")

    info_box(doc,
        "HAI-Prinzipien:  Nur freigegebene Quellen · keine Musterlösungen · Quellenangabe bei jeder Antwort\n"
        "Quiz:  Adaptiv — Schwierigkeit steigt erst, wenn vorherige Stufe sicher beherrscht wird"
    )

    transition_note(doc, "Scrollen zur Kosten / ROI-Sektion")

    # ══════════════════════════════════════════════════════════════════════
    # 11  ROI / KOSTEN  (id="roi")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "11", "Kosten & Nutzen — So entlastet LFA Betrieb & Azubi",
                    "Wirtschaftlicher Nutzen · Drei Kernzahlen · ~2 Minuten")

    stage_direction(doc, "ROI-Sektion sichtbar. Die Zahl '30.000 EUR' ist prominent, grüne Prozentzahlen sichtbar.")

    speaker_block(doc, "WALERI",
        "Kommen wir zur Frage, die Betriebe am Ende immer stellen: Was kostet ein Ausbildungsabbruch wirklich?\n\n"
        "Wir haben das intern durchgerechnet: Ein Azubi, der nach anderthalb Jahren abbricht, "
        "kostet den Betrieb zwischen 42.000 und 49.000 Euro. "
        "Lohn über eineinhalb Jahre, Ausbilderzeit, Anmeldegebühren, und die Kosten für eine neue Stelle — "
        "plus was viele vergessen: Der Azubi wäre nach der Ausbildung ein fertig eingearbeiteter Mitarbeiter geworden.\n\n"
        "LFA greift genau dort an.\n\n"
        "Erstens: 40 Prozent weniger Verwaltungsaufwand für Ausbilder. "
        "Das haben wir nicht geschätzt — das haben wir in unserem eigenen Betrieb gemessen.\n\n"
        "Zweitens: 60 Prozent weniger Zeit, die Azubis außerhalb der Arbeitszeit zum Lernen brauchen, "
        "weil das Lernen gezielter und verständlicher wird. Weniger Frust. Weniger Abbrüche.\n\n"
        "Wer begleitet wird, bleibt. Das ist kein Werbespruch — das ist Erfahrung aus der Praxis.",
        "~2 Min")

    info_box(doc,
        "Kosten im Überblick:\n"
        "  · 42.000 bis 49.000 Euro Kosten pro Ausbildungsabbruch (eigene Kalkulation)\n"
        "  · 40% weniger Verwaltungsaufwand für Ausbilder\n"
        "  · 60% weniger Lernzeit außerhalb der Arbeitszeit\n"
        "  · Ab dem 6. Azubi empfiehlt die IHK einen zweiten Ausbilder — LFA senkt diesen Bedarf"
    )

    transition_note(doc, "Scrollen zur Architektur / Datenschutz-Sektion")

    # ══════════════════════════════════════════════════════════════════════
    # 12  ARCHITEKTUR / TECHNIK  (id="architektur")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "12", "Architektur — Datenschutz & technische Grundlage",
                    "Drei Trust-Karten · ~2 Minuten")

    stage_direction(doc, "Architektur-Sektion sichtbar: drei Vertrauenskarten + Technologie-Laufband.")

    speaker_block(doc, "ELIAS",
        "Ich möchte kurz auf die Grundlage eingehen — nicht um Sie mit Details zu überfordern, "
        "sondern weil Datenschutz in der Ausbildung keine Nebensache ist.\n\n"
        "Drei Punkte:\n\n"
        "Erstens: Die KI erfindet nichts. HAI antwortet nur auf Basis von Inhalten, "
        "die wir ausdrücklich freigegeben haben. Er sucht nicht im Internet. "
        "Was nicht freigegeben ist, existiert für ihn nicht.\n\n"
        "Zweitens: Die Daten aller Nutzer bleiben in Europa. Wir speichern alles auf Servern innerhalb der EU, "
        "nach den Vorgaben der Datenschutzgrundverordnung. Keine Weitergabe an Dritte, "
        "kein Zugriff von außen auf betriebliche Inhalte.\n\n"
        "Drittens: Das ist kein Prototyp. LFA läuft auf bewährten Technologien, "
        "die weltweit täglich von Millionen Menschen genutzt werden. "
        "Das System ist stabil, sicher und kann mit dem Betrieb mitwachsen.",
        "~2 Min")

    info_box(doc,
        "Datenschutz:  DSGVO-konform · EU-Hosting · keine Weitergabe an Dritte\n"
        "KI:  Nur freigegebene Inhalte · kein Internetzugriff · jede Antwort mit Quellenangabe"
    )

    transition_note(doc, "Scrollen zur Team-Sektion")

    # ══════════════════════════════════════════════════════════════════════
    # 13  TEAM  (id="team")
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "13", "Team — Wer steckt dahinter?",
                    "Vorstellung · Waleri · ~1 Minute")

    stage_direction(doc, "Team-Sektion sichtbar. Vier Karten: Waleri Moretz (Gründer), Daniel Moretz (Ausbilder), Leon (Azubi), Elias (Azubi).")

    speaker_block(doc, "WALERI",
        "Ihr habt uns heute schon alle kennengelernt — aber auf der Seite könnt ihr nochmal nachlesen, "
        "wer hinter LFA steckt.\n\n"
        "Daniel hat über neun Jahre IT-Projekterfahrung, ist zertifizierter IHK-Ausbilder "
        "und bildet seit Jahren aktiv aus.\n\n"
        "Leon und Elias sind seit August 2025 bei uns — und haben LFA nicht nur genutzt, "
        "sondern aktiv mitentwickelt. Beide haben eigenständig Funktionen programmiert, "
        "die ihr heute in der Demo gesehen habt. Das ist nicht selbstverständlich nach wenigen Monaten Ausbildung.\n\n"
        "Und ich war selbst mal der Azubi, der nach dem zweiten Schulblock kurz davor war aufzugeben — "
        "trotz guter Noten, weil die Struktur gefehlt hat. "
        "Deshalb gibt es LFA.",
        "~1 Min")

    transition_note(doc, "Scrollen zur FAQ-Sektion — dann Abschluss")

    # ══════════════════════════════════════════════════════════════════════
    # 12  ABSCHLUSS / FAQ-VORBEREITUNG
    # ══════════════════════════════════════════════════════════════════════
    section_heading(doc, "12", "Abschluss & Fragerunde",
                    "Zusammenfassung · Vorbereitete Fragen · ~5–8 Minuten")

    stage_direction(doc, "Optional: kurz durch FAQ-Sektion scrollen und zeigen, dass häufige Fragen dort beantwortet werden.")

    speaker_block(doc, "WALERI",
        "Lassen Sie mich in drei Sätzen zusammenfassen, was Sie heute gesehen haben.\n\n"
        "LFA ist die erste Plattform dieser Art für die IT-Ausbildung in Deutschland — "
        "sie bringt Ausbilder und Azubis gemeinsam in eine strukturierte, auf die IHK abgestimmte Lernumgebung. "
        "Sie nimmt Verwaltungsaufwand weg, macht Lernen greifbar, und gibt jedem Beteiligten "
        "die Klarheit, die er braucht, um die Ausbildung erfolgreich zu beenden.\n\n"
        "Gebaut von WAMOCON in Frankfurt — von einem Team, das diese Ausbildung aus jeder Perspektive kennt: "
        "als Azubi, als Ausbilder, und als Unternehmer.\n\n"
        "Was interessiert Sie am meisten? Was haben Sie noch für Fragen?",
        "~1 Min + Fragen")

    doc.add_paragraph()

    # ── Vorbereitete Fragen ──
    p_faq = doc.add_paragraph()
    r_faq = p_faq.add_run("Vorbereitete Antworten auf IHK-Fragen")
    r_faq.bold = True
    r_faq.font.size = Pt(11)
    r_faq.font.color.rgb = DARK
    p_faq.paragraph_format.space_after = Pt(6)

    faq_items = [
        (
            '\u201eWie stellt ihr sicher, dass die Plattform zur IHK passt?\u201c',
            "Alle Lernbereiche sind direkt auf den offiziellen Ausbildungsrahmenplan der IHK aufgebaut. "
            "Aufgaben, Quizze und das Berichtsheft orientieren sich an den echten IHK-Inhalten. "
            "Das Berichtsheft l\xe4sst sich als PDF exportieren und ist nach den Vorgaben der IHK gegliedert. "
            "Wir sind au\xdferdem im Gespr\xe4ch mit der IHK Frankfurt \xfcber eine engere Zusammenarbeit.",
            "DANIEL"
        ),
        (
            '\u201eWas passiert, wenn die KI etwas Falsches sagt?\u201c',
            "Das ist durch den Aufbau des Systems weitgehend ausgeschlossen: "
            "HAI antwortet nur auf Basis von Inhalten, die wir vorher ausdr\xfccklich freigegeben haben. "
            "Er hat keinen Zugriff auf das Internet. "
            "Jede Antwort zeigt die zugeh\xf6rige Quelle, damit Azubi und Ausbilder nachvollziehen k\xf6nnen, woher sie stammt. "
            "Der Ausbilder hat immer das letzte Wort.",
            "ELIAS"
        ),
        (
            '\u201eNutzt ihr die Plattform schon wirklich?\u201c',
            "Ja \u2014 Leon und Elias nutzen LFA seit dem ersten Tag ihrer Ausbildung. "
            "Das ist kein Testbetrieb, das ist unser echter Alltag. "
            "Wir sind in der Aufbauphase \u2014 das bedeutet: die Kernfunktionen laufen stabil, "
            "aber wir entwickeln weiter. Version 2 ist bereits in Arbeit.",
            "WALERI"
        ),
        (
            '\u201eF\xfcr wen ist LFA gedacht?\u201c',
            "Aktuell f\xfcr Betriebe, die IT-Fachkr\xe4fte in der Fachrichtung Anwendungsentwicklung ausbilden. "
            "Jeder Betrieb bekommt seinen eigenen, vollst\xe4ndig getrennten Bereich in der Plattform. "
            "Wir planen, das Angebot auf weitere IT-Ausbildungsberufe auszuweiten.",
            "DANIEL"
        ),
        (
            '\u201eWas kostet das?\u201c',
            "Die genaue Preisgestaltung stimmen wir gerade ab. "
            "Unser Ziel ist ein monatlicher Beitrag pro Azubi, "
            "der sp\xfcrbar unter den Kosten eines einzigen Ausbildungsabbruchs liegt. "
            "F\xfcr Betriebe, die jetzt in der Aufbauphase einsteigen, planen wir ein besonderes Angebot.",
            "WALERI"
        ),
        (
            '\u201eSind die Daten der Azubis gesch\xfctzt?\u201c',
            "Ja, vollst\xe4ndig. Alle Daten werden auf Servern innerhalb der EU gespeichert, "
            "nach den Vorgaben der Datenschutzgrundverordnung. "
            "Keine Daten werden an Dritte weitergegeben. "
            "Die KI hat keinen Zugriff auf pers\xf6nliche Daten der Nutzer.",
            "ELIAS"
        ),
    ]

    for question, answer, speaker_key in faq_items:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(8)
        p_q.paragraph_format.space_after  = Pt(2)
        rq = p_q.add_run(question)
        rq.bold = True
        rq.italic = True
        rq.font.size = Pt(10)
        rq.font.color.rgb = DARK
        speaker_block(doc, speaker_key, answer)

    # ══════════════════════════════════════════════════════════════════════
    # SCHLUSS-SEITE
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    p_end = doc.add_paragraph()
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_end.paragraph_format.space_before = Pt(80)
    r_end = p_end.add_run("Vielen Dank.")
    r_end.bold = True
    r_end.font.size  = Pt(36)
    r_end.font.color.rgb = RED

    p_end2 = doc.add_paragraph()
    p_end2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_e2 = p_end2.add_run(
        "lfa-learn.com/demo\n"
        "wamocon.com\n\n"
        "WAMOCON GmbH  ·  Frankfurt am Main  ·  2026"
    )
    r_e2.font.size = Pt(11)
    r_e2.font.color.rgb = STAGE

    # ══════════════════════════════════════════════════════════════════════
    # SPEICHERN
    # ══════════════════════════════════════════════════════════════════════
    output_path = r"C:\Users\Leon Moretz\FIAE-Landing-Page\LFA_IHK_Praesentationsskript.docx"
    doc.save(output_path)
    print("OK  Dokument gespeichert: " + output_path)


if __name__ == "__main__":
    build_document()
